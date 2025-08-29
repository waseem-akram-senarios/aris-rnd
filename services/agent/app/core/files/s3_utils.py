import os
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Optional

import boto3
import fitz  # PyMuPDF
from docx import Document  # python-docx


@dataclass
class DocumentContent:
    name: str
    format: str  # 'txt' | 'error' | ext
    bytes_data: bytes


def _filter_filename(filename: str) -> str:
    filtered = Path(filename).stem
    return filtered


def _create_error_document(name: str, message: str) -> DocumentContent:
    return DocumentContent(name=name, format="error", bytes_data=message.encode("utf-8"))


def _extract_docx_text(file_path: str) -> Optional[str]:
    doc = Document(file_path)
    text_parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            text_parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells])
            if row_text.strip():
                text_parts.append(row_text)
    raw = "\n".join(text_parts)
    return raw if raw.strip() else None


def get_document_content_from_s3(bucket: str, key: str) -> DocumentContent:
    s3 = boto3.client("s3")
    doc_name = Path(key).name
    ext = Path(key).suffix.lower()

    # download to /tmp
    safe_name = _filter_filename(doc_name)
    local_path = f"/tmp/{safe_name}{ext if ext else ''}"
    s3.download_file(bucket, key, local_path)

    # PDF: always extract text
    if ext == ".pdf":
        with fitz.open(local_path) as pdf:
            if len(pdf) == 0:
                return _create_error_document(doc_name, "Empty PDF document")
            parts = []
            for i, page in enumerate(pdf):
                page_text = page.get_text()
                if page_text.strip():
                    parts.append(f"[page {i + 1}]\n{page_text}\n")
            if not parts:
                return _create_error_document(doc_name, "PDF has no extractable text")
            content = "".join(parts)
            return DocumentContent(name=doc_name, format="txt", bytes_data=content.encode("utf-8"))

    # DOCX/DOC try extraction
    if ext in {".docx", ".doc"}:
        try:
            text = _extract_docx_text(local_path)
            if text and text.strip():
                return DocumentContent(name=doc_name, format="txt", bytes_data=text.encode("utf-8"))
            return _create_error_document(doc_name, "DOCX has no text content")
        except Exception as exc:
            return _create_error_document(doc_name, f"DOCX extraction failed: {exc}")

    # For other supported types or large files, return raw bytes
    with open(local_path, "rb") as fh:
        data = fh.read()
    return DocumentContent(name=doc_name, format=(ext[1:] if ext else "bin"), bytes_data=data)


